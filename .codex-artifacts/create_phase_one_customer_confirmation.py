from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("docs/customer-confirmation/AI智能下单平台_一期客户确认清单.docx")

# Named compatibility override for Chinese customer-facing output.
LATIN_FONT = "STSong"
CJK_FONT = "STSong"
NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5F6B7A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "D7DEE8"
WHITE = "FFFFFF"


def set_run_font(run, size=11, color="222222", bold=False, italic=False):
    run.font.name = LATIN_FONT
    r_pr = run._element.get_or_add_rPr()
    r_pr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    r_pr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    r_pr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    r_pr.rFonts.set(qn("w:cs"), CJK_FONT)
    r_pr.rFonts.set(qn("w:hint"), "eastAsia")
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")
    lang.set(qn("w:bidi"), "zh-CN")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_style_font(style, size, color="222222", bold=False):
    style.font.name = LATIN_FONT
    r_pr = style._element.get_or_add_rPr()
    r_pr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    r_pr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    r_pr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    r_pr.rFonts.set(qn("w:cs"), CJK_FONT)
    r_pr.rFonts.set(qn("w:hint"), "eastAsia")
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")
    lang.set(qn("w:bidi"), "zh-CN")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_paragraph_border(paragraph, color=BORDER, size="8", space="4", side="bottom"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    tag = OxmlElement(f"w:{side}")
    tag.set(qn("w:val"), "single")
    tag.set(qn("w:sz"), size)
    tag.set(qn("w:space"), space)
    tag.set(qn("w:color"), color)
    p_bdr.append(tag)


def shade_paragraph(paragraph, fill=LIGHT_GRAY):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, value in (("top", 100), ("bottom", 100), ("start", 120), ("end", 120)):
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def add_custom_numbering(doc, marker, font=CJK_FONT):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), marker)
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "269")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "271")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), font)
    fonts.set(qn("w:hAnsi"), font)
    fonts.set(qn("w:eastAsia"), font)
    r_pr.append(fonts)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def add_list_item(doc, text, num_id, bold_prefix=None):
    p = doc.add_paragraph()
    apply_numbering(p, num_id)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_body(doc, text, bold=False, color="222222", after=6, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, bold=bold, color=color, italic=italic)
    return p


def add_label_paragraph(doc, label, text, fill=LIGHT_GRAY):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.25
    shade_paragraph(p, fill)
    add_paragraph_border(p, color=BLUE, size="18", space="5", side="left")
    r1 = p.add_run(label)
    set_run_font(r1, bold=True, color=NAVY)
    r2 = p.add_run(text)
    set_run_font(r2, color="26323F")
    return p


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    normal = doc.styles["Normal"]
    set_style_font(normal, 11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, 16, BLUE, True)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, 13, BLUE, True)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, 12, DARK_BLUE, True)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    # Quiet customer-pack running header.
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(3)
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    r1 = hp.add_run("AI 智能下单平台")
    set_run_font(r1, size=8.5, color=MUTED, bold=True)
    r2 = hp.add_run("\t一期客户确认清单")
    set_run_font(r2, size=8.5, color=MUTED)
    add_paragraph_border(hp, color=BORDER, size="6", space="3", side="bottom")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(2)
    run = fp.add_run("第 ")
    set_run_font(run, size=8.5, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    tail = fp.add_run(" 页")
    set_run_font(tail, size=8.5, color=MUTED)

    return doc


def build_document():
    doc = setup_document()
    bullet_id = add_custom_numbering(doc, "•")
    check_id = add_custom_numbering(doc, "☐")

    # Customer-pack title block.
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(2)
    kr = kicker.add_run("客户确认资料｜一期交付")
    set_run_font(kr, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    tr = title.add_run("AI 智能下单平台\n一期客户确认清单")
    set_run_font(tr, size=27, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    sr = subtitle.add_run("客户确认事项、业务资料输入与一期最终验收配合清单")
    set_run_font(sr, size=12.5, color=MUTED)

    meta = doc.add_table(rows=2, cols=2)
    meta.style = "Table Grid"
    set_table_geometry(meta, [4680, 4680])
    values = [
        ("文档状态", "待客户确认"),
        ("适用阶段", "一期交付收口"),
        ("确认方式", "仅确认关键事项，不逐功能签字"),
        ("版本日期", "2026 年 7 月 20 日"),
    ]
    for idx, (label, value) in enumerate(values):
        cell = meta.rows[idx // 2].cells[idx % 2]
        set_cell_shading(cell, LIGHT_BLUE if idx in (0, 1) else WHITE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r1 = p.add_run(label + "\n")
        set_run_font(r1, size=8.5, color=MUTED, bold=True)
        r2 = p.add_run(value)
        set_run_font(r2, size=10.5, color=NAVY, bold=True)
    set_repeat_table_header(meta.rows[0])

    add_label_paragraph(
        doc,
        "结论：",
        "按当前项目确认口径，客户不需要对 38 个功能逐项签字。当前需要客户/PM拍板 2 项，另需提供 2 类业务资料；功能、真实环境和培训全部完成后，再签署一份一期总体验收记录。",
        LIGHT_GRAY,
    )

    doc.add_heading("项目进度概览（截至 2026 年 7 月 20 日）", level=1)
    add_body(doc, "AI 智能下单平台一期核心功能开发已基本完成，目前已进入真实环境联调及最终交付验收阶段。", bold=True, color=NAVY)
    add_body(doc, "一期已完成医生端、客服端、生产端和管理端四个业务端口。医生下单、客服审核、生产流转、设计稿确认、入检 / 出检、返工、工时绩效、账单物流、通知及 AI 辅助等核心流程已形成完整闭环。")
    add_body(doc, "系统已通过 12 步主业务链路验证、四端页面验收及后端回归测试，并已具备独立演示环境和客户演示条件。")
    add_body(doc, "按 PRD V2 原 38 项验收标准统计，目前 30 项已通过，1 项待补充标准工时数据后复核，7 项已完成系统能力建设、待真实模型、对象存储、网络及部署环境下进行验证，当前不存在明确的本地功能缺失项。")
    add_body(doc, "下一阶段重点工作包括：", bold=True, color=DARK_BLUE)
    for item in (
        "确认各产品动态表单最终字段及文件上传参数；",
        "补充生产备注模板和各工序标准工时数据；",
        "完成真实 AI 模型、对象存储、HTTPS、通知及生产环境联调；",
        "完成备份恢复、监控告警、发布回滚、客户培训及总体验收。",
    ):
        add_list_item(doc, item, bullet_id)
    add_label_paragraph(doc, "综合判断：", "一期已完成主体建设，当前处于上线前最后收口阶段。待真实环境联调和最终验收完成后，即可进入正式上线交付。", LIGHT_BLUE)

    doc.add_heading("一、现在需要客户 / PM 确认", level=1)

    doc.add_heading("1. 动态下单表单最终版", level=2)
    add_body(doc, "按不同产品确认医生下单时需要填写哪些内容、哪些内容必填、有哪些固定选项，以及各项内容的显示顺序。")
    add_label_paragraph(doc, "建议方式：", "以当前系统内容为底稿，让客户只标注保留、删除或修改，不要求客户从零设计表单。", LIGHT_BLUE)

    doc.add_heading("2. 文件上传限制", level=2)
    add_body(doc, "请确认是否采用当前一期默认值：")
    for item in (
        "单个文件最大 500MB；",
        "每个订单最多上传 50 个文件；",
        "允许上传 STL、SLA、PDF、图片、TXT、ZIP 等常用文件。",
    ):
        add_list_item(doc, item, bullet_id)
    add_body(doc, "如客户没有特殊要求，可直接采用上述默认值。", bold=True, color=DARK_BLUE)

    doc.add_heading("二、客户或工厂需要提供", level=1)

    doc.add_heading("1. AI 生产备注正式模板", level=2)
    add_body(doc, "请提供工厂日常使用的生产备注格式、固定栏目和示例。AI 只生成草稿，仍由客服或生产人员人工确认后使用。")
    add_label_paragraph(doc, "没有固定模板时：", "可提供几份已经脱敏的历史生产备注、生产单或客服整理后的制作要求，由项目组提炼统一格式后再交客户确认。", LIGHT_BLUE)

    doc.add_heading("2. 各工序标准工时表", level=2)
    add_body(doc, "请由工厂业务负责人提供每道工序的标准时间，用于核算工时效率、准时率等绩效指标。")
    add_body(doc, "这属于业务数据输入，不是产品需求确认。项目方需要先指定一位负责人，协调工厂业务人员收集。")

    doc.add_heading("三、完成一期前需要客户配合验收", level=1)

    doc.add_heading("1. 按真实业务走一遍完整流程", level=2)
    flow = add_body(doc, "医生下单 → 客服审核 → 生产审核 → 派工 → 入检 → 生产 → 出检 / 返工 → 终检 → 录入物流 → 医生确认收货。", bold=True, color=NAVY, after=10)
    shade_paragraph(flow, LIGHT_BLUE)

    doc.add_heading("2. 验收关键异常情况", level=2)
    for item in (
        "资料不全时退回补充；",
        "设计稿驳回后重新上传；",
        "生产出检不通过后进入返工；",
        "无权限人员不能查看文件；",
        "不同诊所之间不能互相查看订单。",
    ):
        add_list_item(doc, item, bullet_id)

    doc.add_heading("3. 在真实环境验收 AI", level=2)
    add_body(doc, "重点确认翻译是否准确、客服查询是否可靠、医生端不会看到工厂内部信息，以及生产备注是否符合客户提供的模板。")

    doc.add_heading("4. 在真实环境验收文件上传", level=2)
    add_body(doc, "验证大文件上传、网络中断后恢复、换设备继续上传、文件预览以及访问权限限制。")

    doc.add_heading("5. 完成上线环境验收", level=2)
    add_body(doc, "包括正式域名和安全访问、数据备份恢复、系统监控及故障回退。主要由部署团队执行，客户确认结果。")

    doc.add_heading("6. 完成四端培训", level=2)
    add_body(doc, "分别完成医生端、客服端、生产端和管理端培训，并记录参会人员、培训问题和培训结论。")

    doc.add_heading("7. 签署一期总体验收记录", level=2)
    add_body(doc, "不需要对 38 项功能分别签字。所有功能、真实环境和培训完成后，由客户代表与 PM 签署一份一期总体验收记录即可。", bold=True, color=NAVY)

    doc.add_heading("四、当前项目状态与近期重点", level=1)
    add_body(doc, "根据当前项目记录，原 PRD 的 38 项验收中：")

    status_table = doc.add_table(rows=1, cols=3)
    status_table.style = "Table Grid"
    set_table_geometry(status_table, [2900, 2900, 3560])
    headers = ["当前状态", "数量", "说明"]
    for idx, text in enumerate(headers):
        cell = status_table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=10, color=NAVY, bold=True)
    set_repeat_table_header(status_table.rows[0])
    rows = [
        ("已有本地通过证据", "30 项", "本地代码、测试或演示链路已有证据"),
        ("仍待业务数据", "1 项", "缺少各工序标准工时"),
        ("仍待真实条件验收", "7 项", "需要真实 AI、真实网络或真实环境验证"),
    ]
    for ridx, row_values in enumerate(rows):
        cells = status_table.add_row().cells
        for idx, value in enumerate(row_values):
            set_cell_width(cells[idx], [2900, 2900, 3560][idx])
            if ridx % 2 == 1:
                set_cell_shading(cells[idx], "FAFBFC")
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx < 2 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=10, color=NAVY if idx < 2 else "26323F", bold=idx < 2)

    add_label_paragraph(
        doc,
        "近期最应优先向客户推进：",
        "动态下单表单、文件上传限制、AI 生产备注模板和各工序标准工时表。",
        LIGHT_BLUE,
    )

    doc.add_heading("五、一期不作为必须完成项", level=1)
    for item in (
        "真实支付平台接入；",
        "物流平台自动接口；",
        "真实电子签章。",
    ):
        add_list_item(doc, item, bullet_id)
    add_body(doc, "如后续新增以上要求，应作为范围变更单独评估。", color=MUTED)

    doc.add_heading("六、客户确认与配合清单", level=1)
    checklist = [
        "确认动态下单表单最终内容；",
        "确认文件大小、数量和允许上传的文件类型；",
        "提供 AI 生产备注正式模板或脱敏历史样例；",
        "提供各工序标准工时表；",
        "功能完成后配合完整业务流程验收；",
        "配合真实环境 AI 和文件上传验收；",
        "配合四端培训；",
        "最终签署一份一期总体验收记录。",
    ]
    for item in checklist:
        add_list_item(doc, item, check_id)

    closing = add_body(doc, "本清单用于一期客户确认与交付准备，不代表尚未发生的真实环境验收、培训或最终签收已经完成。", bold=True, color=DARK_BLUE, after=0)
    shade_paragraph(closing, LIGHT_GRAY)

    # Core properties and save.
    doc.core_properties.title = "AI 智能下单平台一期客户确认清单"
    doc.core_properties.subject = "一期客户确认事项、业务资料输入与最终验收配合"
    doc.core_properties.author = "AI 智能下单平台项目组"
    doc.core_properties.keywords = "一期, 客户确认, 验收, 动态表单, 标准工时"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build_document()
