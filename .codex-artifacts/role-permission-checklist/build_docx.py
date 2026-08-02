#!/usr/bin/env python3
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/yuri/Documents/AI智能下单平台")
SOURCE = ROOT / "docs/customer-confirmation/AI智能下单平台_四端角色职责与权限客户确认清单_20260724.md"
OUTPUT = ROOT / "docs/customer-confirmation/AI智能下单平台_四端角色职责与权限客户确认清单_20260724.docx"
SKILL_ROOT = Path("/Users/yuri/.codex/plugins/cache/openai-primary-runtime/documents/26.715.12143/skills/documents")
sys.path.insert(0, str(SKILL_ROOT / "scripts"))

from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


# compact_reference_guide preset, with named CJK and dense-matrix overrides.
CONTENT_WIDTH_DXA = 9360
METADATA_TABLE_INDENT_DXA = 120
CALLOUT_TABLE_INDENT_DXA = 180
CONTENT_TABLE_INDENT_DXA = 100
FONT_LATIN = "Arial Unicode MS"
FONT_CJK = "Arial Unicode MS"
NAVY = "17324D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "5B6472"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
RESPONSE_FILL = "F8FAFC"
BORDER = "CBD5E1"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = BORDER, size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_run_font(run, size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT_LATIN
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT_LATIN)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT_LATIN)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_CJK)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_inline_runs(paragraph, text: str, *, size: float | None = None, color: str | None = None) -> None:
    # Minimal Markdown inline support for bold and code markers.
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size, color=DARK_BLUE, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    doc.styles["Heading 1"].paragraph_format.page_break_before = True

    response = doc.styles.add_style("Response Line", WD_STYLE_TYPE.PARAGRAPH)
    response.base_style = normal
    response.font.name = FONT_LATIN
    response._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    response.font.size = Pt(10.5)
    response.font.color.rgb = RGBColor.from_string(MUTED)
    response.font.italic = True
    response.paragraph_format.space_before = Pt(1)
    response.paragraph_format.space_after = Pt(7)
    response.paragraph_format.keep_together = True

    table_text = doc.styles.add_style("Compact Table Text", WD_STYLE_TYPE.PARAGRAPH)
    table_text.base_style = normal
    table_text.font.name = FONT_LATIN
    table_text._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    table_text.font.size = Pt(9)
    table_text.paragraph_format.space_before = Pt(0)
    table_text.paragraph_format.space_after = Pt(0)
    table_text.paragraph_format.line_spacing = 1.08


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    field_run = paragraph.add_run()
    set_run_font(field_run, size=9, color=MUTED)
    field_run._r.append(fld_char_begin)
    field_run._r.append(instr_text)
    field_run._r.append(fld_char_end)
    end = paragraph.add_run(" 页")
    set_run_font(end, size=9, color=MUTED)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("AI 智能下单平台｜角色职责与权限确认")
    set_run_font(run, size=9, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    add_page_field(fp)


def add_cover(doc: Document, callout_text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("客户需求确认文件")
    set_run_font(run, size=10.5, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("AI 智能下单平台")
    set_run_font(run, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("四端角色职责与权限客户确认清单")
    set_run_font(run, size=16, color=BLUE, bold=True)

    metadata = doc.add_table(rows=3, cols=2)
    metadata_data = [
        ("版本", "客户确认草案 V1.0"),
        ("日期", "2026-07-24"),
        ("适用范围", "医生端、客服端、生产端、管理端"),
    ]
    for row, (label, value) in zip(metadata.rows, metadata_data):
        row.cells[0].text = ""
        row.cells[1].text = ""
        lp = row.cells[0].paragraphs[0]
        vp = row.cells[1].paragraphs[0]
        add_inline_runs(lp, label, size=10, color=MUTED)
        add_inline_runs(vp, value, size=10.5, color=INK)
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], LIGHT_FILL)
        for cell in row.cells:
            set_cell_borders(cell, color=BORDER, size="3")
    apply_table_geometry(metadata, [1700, 7660], table_width_dxa=CONTENT_WIDTH_DXA, indent_dxa=METADATA_TABLE_INDENT_DXA)

    doc.add_paragraph()
    callout = doc.add_table(rows=1, cols=1)
    cell = callout.cell(0, 0)
    cell.text = ""
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_after = Pt(0)
    add_inline_runs(cp, callout_text, size=10.5, color=NAVY)
    set_cell_shading(cell, HEADER_FILL)
    set_cell_borders(cell, color="B9C9D8", size="6")
    apply_table_geometry(callout, [CONTENT_WIDTH_DXA], table_width_dxa=CONTENT_WIDTH_DXA, indent_dxa=CALLOUT_TABLE_INDENT_DXA, cell_margins_dxa={"top": 140, "bottom": 140, "start": 180, "end": 180})

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("建议确认顺序")
    set_run_font(run, size=12, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("四端角色 → 数据范围与关键动作 → 跨端主责与复核 → 最终角色名称")
    set_run_font(run, size=10.2, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("填写标记：□保留　□合并　□取消　□新增　｜　空白横线用于填写客户意见")
    set_run_font(run, size=9.5, color=MUTED, italic=True)
    doc.add_page_break()


def add_numbering_definition(doc: Document, *, start: int, bullet: bool = False, level: int = 0) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start_el = OxmlElement("w:start")
    start_el.set(qn("w:val"), str(start))
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), ("•" if level == 0 else "◦") if bullet else "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    marker = 270 + level * 360
    text_indent = 540 + level * 360
    tab.set(qn("w:pos"), str(text_indent))
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(text_indent))
    ind.set(qn("w:hanging"), str(text_indent - marker))
    p_pr.append(tabs)
    p_pr.append(ind)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT_LATIN)
    fonts.set(qn("w:hAnsi"), FONT_LATIN)
    fonts.set(qn("w:eastAsia"), FONT_CJK)
    r_pr.append(fonts)
    lvl.extend([start_el, num_fmt, lvl_text, suff, p_pr, r_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


_bullet_num_ids: dict[int, int] = {}


def add_bullet_paragraph(doc: Document, text: str, *, level: int = 0):
    if level not in _bullet_num_ids:
        _bullet_num_ids[level] = add_numbering_definition(doc, start=1, bullet=True, level=level)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    apply_num(p, _bullet_num_ids[level])
    add_inline_runs(p, text, size=10.7, color=INK)
    return p


def add_numbered_paragraph(doc: Document, text: str, number: int):
    num_id = add_numbering_definition(doc, start=number, bullet=False)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_with_next = True
    apply_num(p, num_id)
    add_inline_runs(p, text, size=10.8, color=INK)
    return p


def table_weights(rows: list[list[str]]) -> list[float]:
    col_count = len(rows[0])
    max_lengths = []
    for idx in range(col_count):
        max_len = max(len(row[idx]) for row in rows)
        max_lengths.append(max_len)
    weights = [max(0.85, min(3.2, length ** 0.52)) for length in max_lengths]
    if col_count >= 5:
        weights[0] = max(weights[0], 1.6)
    return weights


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.allow_autofit = False
    font_size = 9.6 if col_count <= 3 else 8.8 if col_count == 4 else 8.1

    for row_idx, source_row in enumerate(rows):
        table_row = table.rows[row_idx]
        for col_idx, text in enumerate(source_row):
            cell = table_row.cells[col_idx]
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.style = doc.styles["Compact Table Text"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == col_count - 1 and len(text) < 22 else WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(p, text.replace("<br>", "\n"), size=font_size, color=NAVY if row_idx == 0 else INK)
            for run in p.runs:
                if row_idx == 0 or (col_idx == 0 and len(text) < 24):
                    run.bold = True
            set_cell_shading(cell, HEADER_FILL if row_idx == 0 else WHITE)
            set_cell_borders(cell)
        if row_idx == 0:
            set_repeat_table_header(table_row)

    widths = column_widths_from_weights(table_weights(rows), CONTENT_WIDTH_DXA)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=CONTENT_TABLE_INDENT_DXA,
        cell_margins_dxa={"top": 90, "bottom": 90, "start": 100, "end": 100},
    )
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)


def parse_table(lines: list[str], start_index: int) -> tuple[list[list[str]], int]:
    raw = []
    i = start_index
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw.append(lines[i].strip())
        i += 1
    rows = []
    for idx, line in enumerate(raw):
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if idx == 1 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows, i


def is_response_line(text: str) -> bool:
    prefixes = (
        "客户意见", "客户规则", "客户确认", "客户修改", "客户补充", "客户实际",
        "角色处理", "客户单位", "确认人及职务", "确认日期", "四端最终角色名称",
        "仍待确认事项", "请确认是否", "客户是否需要限制", "生产端：", "医生端：",
        "客服端：", "管理端：", "内返：", "外返：",
    )
    return text.startswith(prefixes) or bool(re.fullmatch(r"_+", text.replace(" ", "")))


def add_response_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Response Line")
    add_inline_runs(p, text, size=10.5, color=MUTED)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), RESPONSE_FILL)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), BORDER)
    borders.append(bottom)
    p_pr.append(borders)


def add_divider(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D8DEE7")
    borders.append(bottom)
    p_pr.append(borders)


def build() -> None:
    source_text = SOURCE.read_text(encoding="utf-8")
    lines = source_text.splitlines()
    callout = next(line[2:].strip() for line in lines if line.startswith("> "))

    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_cover(doc, callout)

    # Start the main body at the first numbered section; the cover replaces source title metadata.
    start = next(i for i, line in enumerate(lines) if line.startswith("## 一、"))
    i = start
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|"):
            table_rows, i = parse_table(lines, i)
            add_markdown_table(doc, table_rows)
            continue
        if stripped == "---":
            i += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(stripped[4:].strip(), style="Heading 2")
            for run in p.runs:
                set_run_font(run, size=13, color=BLUE, bold=True)
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph(stripped[3:].strip(), style="Heading 1")
            for run in p.runs:
                set_run_font(run, size=16, color=BLUE, bold=True)
            i += 1
            continue
        bullet_match = re.match(r"^(\s*)-\s+(.*)$", raw)
        if bullet_match:
            indent = len(bullet_match.group(1))
            add_bullet_paragraph(doc, bullet_match.group(2).strip(), level=1 if indent >= 2 else 0)
            i += 1
            continue
        number_match = re.match(r"^\s*(\d+)\.\s+(.*)$", raw)
        if number_match:
            add_numbered_paragraph(doc, number_match.group(2).strip(), int(number_match.group(1)))
            i += 1
            continue
        if stripped.startswith("> "):
            i += 1
            continue
        if is_response_line(stripped):
            add_response_paragraph(doc, stripped)
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.keep_together = False
        add_inline_runs(p, stripped, size=10.8, color=INK)
        i += 1

    doc.core_properties.title = "AI智能下单平台_四端角色职责与权限客户确认清单"
    doc.core_properties.subject = "医生端、客服端、生产端、管理端角色职责与权限确认"
    doc.core_properties.author = "AI智能下单平台项目组"
    doc.core_properties.keywords = "角色, 权限, 客户确认, 医生端, 客服端, 生产端, 管理端"
    doc.core_properties.comments = "客户确认草案 V1.0"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
