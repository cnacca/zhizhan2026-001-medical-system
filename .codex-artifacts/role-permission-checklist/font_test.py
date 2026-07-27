#!/usr/bin/env python3
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

out = Path("/Users/yuri/Documents/AI智能下单平台/.codex-artifacts/role-permission-checklist/font-test.docx")
fonts = [
    "Arial Unicode MS",
    "Hiragino Sans GB",
    "Heiti SC",
    "STHeiti",
    "Songti SC",
    "STSong",
    "SimSong",
    "Kaiti SC",
    "PingFang SC",
    "Apple LiGothic",
    "Apple LiSung",
]
doc = Document()
for name in fonts:
    p = doc.add_paragraph()
    run = p.add_run(f"{name}：中文字体测试，医生端、客服端、生产端、管理端。")
    run.font.name = name
    run.font.size = Pt(16)
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), name)
doc.save(out)
print(out)
